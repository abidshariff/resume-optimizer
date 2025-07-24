"""
Enhanced Word document generation using python-docx library.
Creates a professionally formatted resume in .docx format.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_enhanced_word_resume(resume_json):
    """
    Create a professionally formatted Word document using python-docx.
    
    Args:
        resume_json (dict): Resume data
    
    Returns:
        bytes: Generated Word document as bytes
    """
    
    try:
        # Create a new Document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Define styles
        styles = doc.styles
        
        # Name style
        name_style = styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
        name_font = name_style.font
        name_font.name = 'Calibri'
        name_font.size = Pt(18)
        name_font.bold = True
        name_font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dark blue
        
        # Contact style
        contact_style = styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
        contact_font = contact_style.font
        contact_font.name = 'Calibri'
        contact_font.size = Pt(11)
        
        # Section header style
        header_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        header_font = header_style.font
        header_font.name = 'Calibri'
        header_font.size = Pt(14)
        header_font.bold = True
        header_font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dark blue
        
        # Job title style
        job_title_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
        job_title_font = job_title_style.font
        job_title_font.name = 'Calibri'
        job_title_font.size = Pt(12)
        job_title_font.bold = True
        
        # Job details style
        job_details_style = styles.add_style('JobDetails', WD_STYLE_TYPE.PARAGRAPH)
        job_details_font = job_details_style.font
        job_details_font.name = 'Calibri'
        job_details_font.size = Pt(11)
        job_details_font.italic = True
        
        # Normal text style
        normal_style = styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
        # Add name
        name = resume_json.get('full_name', 'Your Name')
        name_paragraph = doc.add_paragraph(name, style='Name')
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info
        contact = resume_json.get('contact_info', 'Your Contact Information')
        contact_paragraph = doc.add_paragraph(contact, style='Contact')
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line
        add_horizontal_line(doc)
        
        # Add Professional Summary section
        doc.add_paragraph('PROFESSIONAL SUMMARY', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        summary = resume_json.get('professional_summary', 'Professional summary not available')
        doc.add_paragraph(summary, style='NormalText')
        
        # Add Skills section
        doc.add_paragraph('SKILLS', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        skills = resume_json.get('skills', [])
        if skills:
            skills_paragraph = doc.add_paragraph(style='NormalText')
            for i, skill in enumerate(skills):
                if i > 0:
                    skills_paragraph.add_run(' • ')
                skills_paragraph.add_run(skill)
        
        # Add Experience section
        doc.add_paragraph('EXPERIENCE', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        experience = resume_json.get('experience', [])
        for job in experience:
            # Job title
            title = job.get('title', 'Job Title')
            doc.add_paragraph(title, style='JobTitle')
            
            # Company and dates
            company = job.get('company', 'Company')
            dates = job.get('dates', 'Dates')
            doc.add_paragraph(f"{company} | {dates}", style='JobDetails')
            
            # Achievements
            achievements = job.get('achievements', [])
            for achievement in achievements:
                achievement_para = doc.add_paragraph(style='NormalText')
                achievement_para.add_run('• ').bold = True
                achievement_para.add_run(achievement)
            
            # Add space between jobs
            doc.add_paragraph()
        
        # Add Education section
        doc.add_paragraph('EDUCATION', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        education = resume_json.get('education', [])
        for edu in education:
            # Degree
            degree = edu.get('degree', 'Degree')
            doc.add_paragraph(degree, style='JobTitle')
            
            # Institution and dates
            institution = edu.get('institution', 'Institution')
            dates = edu.get('dates', 'Dates')
            doc.add_paragraph(f"{institution} | {dates}", style='JobDetails')
            
            # Details if available
            if 'details' in edu and edu['details']:
                doc.add_paragraph(edu['details'], style='NormalText')
        
        # Save the document to a bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error creating enhanced Word document: {str(e)}")
        raise e

def add_horizontal_line(doc, thin=False):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    
    # Create the line
    run = p.add_run()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6' if thin else '12')  # 6 for thin line, 12 for thick line
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '1F4E79')  # Dark blue
    border.append(bottom)
    
    # Add the border to the paragraph
    p._p.get_or_add_pPr().append(border)
