#!/usr/bin/env python3
"""
Script to create a professional one-page resume template for docxtpl.
This creates a Word document with placeholders that can be filled by python-docx-template.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_resume_template():
    """Create a professional one-page resume template with docxtpl placeholders."""
    
    # Create document
    doc = Document()
    
    # Set document margins (narrow margins for one-page layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Define styles
    styles = doc.styles
    
    # Header style for name
    if 'Name Header' not in [s.name for s in styles]:
        name_style = styles.add_style('Name Header', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Calibri'
        name_style.font.size = Pt(24)
        name_style.font.bold = True
        name_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Professional blue
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_style.paragraph_format.space_after = Pt(6)
    
    # Contact info style
    if 'Contact Info' not in [s.name for s in styles]:
        contact_style = styles.add_style('Contact Info', WD_STYLE_TYPE.PARAGRAPH)
        contact_style.font.name = 'Calibri'
        contact_style.font.size = Pt(11)
        contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_style.paragraph_format.space_after = Pt(12)
    
    # Section header style
    if 'Section Header' not in [s.name for s in styles]:
        section_style = styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Calibri'
        section_style.font.size = Pt(14)
        section_style.font.bold = True
        section_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        section_style.paragraph_format.space_before = Pt(12)
        section_style.paragraph_format.space_after = Pt(6)
        # Add bottom border
        section_style.paragraph_format.keep_with_next = True
    
    # Job title style
    if 'Job Title' not in [s.name for s in styles]:
        job_title_style = styles.add_style('Job Title', WD_STYLE_TYPE.PARAGRAPH)
        job_title_style.font.name = 'Calibri'
        job_title_style.font.size = Pt(12)
        job_title_style.font.bold = True
        job_title_style.paragraph_format.space_after = Pt(3)
    
    # Company info style
    if 'Company Info' not in [s.name for s in styles]:
        company_style = styles.add_style('Company Info', WD_STYLE_TYPE.PARAGRAPH)
        company_style.font.name = 'Calibri'
        company_style.font.size = Pt(11)
        company_style.font.italic = True
        company_style.paragraph_format.space_after = Pt(6)
    
    # Body text style
    if 'Body Text' not in [s.name for s in styles]:
        body_style = styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Calibri'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Bullet style
    if 'Bullet Text' not in [s.name for s in styles]:
        bullet_style = styles.add_style('Bullet Text', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Calibri'
        bullet_style.font.size = Pt(10)
        bullet_style.paragraph_format.space_after = Pt(3)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
    
    # === HEADER SECTION ===
    # Name
    name_para = doc.add_paragraph('{{full_name}}', style='Name Header')
    
    # Contact Information
    contact_para = doc.add_paragraph('{{contact_info}}', style='Contact Info')
    
    # === PROFESSIONAL SUMMARY ===
    summary_header = doc.add_paragraph('PROFESSIONAL SUMMARY', style='Section Header')
    # Add underline to section headers
    summary_header.runs[0].font.underline = True
    
    summary_para = doc.add_paragraph('{{professional_summary}}', style='Body Text')
    
    # === SKILLS ===
    skills_header = doc.add_paragraph('CORE COMPETENCIES', style='Section Header')
    skills_header.runs[0].font.underline = True
    
    # Skills will be displayed as a formatted list
    skills_para = doc.add_paragraph(style='Body Text')
    skills_para.add_run('{% for skill in skills %}{{skill}}{% if not loop.last %} • {% endif %}{% endfor %}')
    
    # === PROFESSIONAL EXPERIENCE ===
    exp_header = doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Section Header')
    exp_header.runs[0].font.underline = True
    
    # Experience loop
    exp_loop_start = doc.add_paragraph('{% for job in experience %}', style='Body Text')
    exp_loop_start.runs[0].font.size = Pt(1)  # Make template code invisible
    exp_loop_start.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Job title and company
    job_title_para = doc.add_paragraph('{{job.title}}', style='Job Title')
    company_para = doc.add_paragraph('{{job.company}} | {{job.dates}}', style='Company Info')
    
    # Job responsibilities
    responsibilities_para = doc.add_paragraph('{% for responsibility in job.responsibilities %}', style='Bullet Text')
    responsibilities_para.runs[0].font.size = Pt(1)
    responsibilities_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    bullet_para = doc.add_paragraph('• {{responsibility}}', style='Bullet Text')
    
    end_resp_para = doc.add_paragraph('{% endfor %}', style='Bullet Text')
    end_resp_para.runs[0].font.size = Pt(1)
    end_resp_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    exp_loop_end = doc.add_paragraph('{% endfor %}', style='Body Text')
    exp_loop_end.runs[0].font.size = Pt(1)
    exp_loop_end.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # === EDUCATION ===
    edu_header = doc.add_paragraph('EDUCATION', style='Section Header')
    edu_header.runs[0].font.underline = True
    
    # Education loop
    edu_loop_start = doc.add_paragraph('{% for edu in education %}', style='Body Text')
    edu_loop_start.runs[0].font.size = Pt(1)
    edu_loop_start.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    degree_para = doc.add_paragraph('{{edu.degree}}', style='Job Title')
    institution_para = doc.add_paragraph('{{edu.institution}} | {{edu.dates}}', style='Company Info')
    
    edu_loop_end = doc.add_paragraph('{% endfor %}', style='Body Text')
    edu_loop_end.runs[0].font.size = Pt(1)
    edu_loop_end.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    return doc

def main():
    """Create and save the resume template."""
    doc = create_resume_template()
    doc.save('resume_template.docx')
    print("Resume template created: resume_template.docx")
    print("\nTemplate placeholders:")
    print("- {{full_name}}")
    print("- {{contact_info}}")
    print("- {{professional_summary}}")
    print("- {{skills}} (array)")
    print("- {{experience}} (array with title, company, dates, responsibilities)")
    print("- {{education}} (array with degree, institution, dates)")

if __name__ == '__main__':
    main()
