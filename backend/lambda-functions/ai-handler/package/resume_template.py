"""
Resume Template Generator using python-docx
Creates professional Word documents with predefined formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import json
import re
from datetime import datetime

class ResumeTemplate:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for the resume"""
        styles = self.doc.styles
        
        # Header style for name
        if 'Resume Name' not in [s.name for s in styles]:
            name_style = styles.add_style('Resume Name', WD_STYLE_TYPE.PARAGRAPH)
            name_font = name_style.font
            name_font.name = 'Calibri'
            name_font.size = Pt(24)
            name_font.bold = True
            name_font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Professional blue
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_style.paragraph_format.space_after = Pt(6)
        
        # Contact info style
        if 'Contact Info' not in [s.name for s in styles]:
            contact_style = styles.add_style('Contact Info', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Calibri'
            contact_font.size = Pt(11)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(12)
        
        # Section header style
        if 'Section Header' not in [s.name for s in styles]:
            section_style = styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(14)
            section_font.bold = True
            section_font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
            
            # Add bottom border to section headers
            p_format = section_style.paragraph_format
            p_border = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), '2F5496')
            p_border.append(bottom_border)
            section_style._element.get_or_add_pPr().append(p_border)
        
        # Job title style
        if 'Job Title' not in [s.name for s in styles]:
            job_style = styles.add_style('Job Title', WD_STYLE_TYPE.PARAGRAPH)
            job_font = job_style.font
            job_font.name = 'Calibri'
            job_font.size = Pt(12)
            job_font.bold = True
            job_style.paragraph_format.space_after = Pt(3)
        
        # Company/Date style
        if 'Company Date' not in [s.name for s in styles]:
            company_style = styles.add_style('Company Date', WD_STYLE_TYPE.PARAGRAPH)
            company_font = company_style.font
            company_font.name = 'Calibri'
            company_font.size = Pt(11)
            company_font.italic = True
            company_style.paragraph_format.space_after = Pt(6)
        
        # Bullet point style
        if 'Bullet Point' not in [s.name for s in styles]:
            bullet_style = styles.add_style('Bullet Point', WD_STYLE_TYPE.PARAGRAPH)
            bullet_font = bullet_style.font
            bullet_font.name = 'Calibri'
            bullet_font.size = Pt(11)
            bullet_style.paragraph_format.left_indent = Inches(0.25)
            bullet_style.paragraph_format.space_after = Pt(3)
    
    def parse_resume_content(self, resume_text):
        """Parse the AI-generated resume text into structured data"""
        try:
            # Try to parse as JSON first
            if resume_text.strip().startswith('{'):
                return json.loads(resume_text)
        except:
            pass
        
        # Parse plain text resume
        sections = {}
        current_section = None
        current_content = []
        
        lines = resume_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this is a section header (all caps or title case with colons)
            if (line.isupper() or line.endswith(':')) and len(line) < 50:
                if current_section:
                    sections[current_section] = '\n'.join(current_content)
                current_section = line.replace(':', '').strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Add the last section
        if current_section:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def add_header(self, name, contact_info):
        """Add the header with name and contact information"""
        # Add name
        name_para = self.doc.add_paragraph(name, style='Resume Name')
        
        # Add contact info
        if isinstance(contact_info, dict):
            contact_parts = []
            if contact_info.get('email'):
                contact_parts.append(contact_info['email'])
            if contact_info.get('phone'):
                contact_parts.append(contact_info['phone'])
            if contact_info.get('location'):
                contact_parts.append(contact_info['location'])
            if contact_info.get('linkedin'):
                contact_parts.append(contact_info['linkedin'])
            
            contact_text = ' | '.join(contact_parts)
        else:
            contact_text = str(contact_info)
        
        if contact_text:
            self.doc.add_paragraph(contact_text, style='Contact Info')
    
    def add_section(self, title, content):
        """Add a section with title and content"""
        # Add section header
        self.doc.add_paragraph(title.upper(), style='Section Header')
        
        if isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    self._add_experience_item(item)
                else:
                    self.doc.add_paragraph(f"• {item}", style='Bullet Point')
        elif isinstance(content, dict):
            if 'items' in content:
                for item in content['items']:
                    self._add_experience_item(item)
            else:
                self._add_experience_item(content)
        else:
            # Handle plain text content
            paragraphs = content.split('\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    if para.startswith('•') or para.startswith('-'):
                        self.doc.add_paragraph(para, style='Bullet Point')
                    else:
                        self.doc.add_paragraph(para)
    
    def _add_experience_item(self, item):
        """Add an experience/education item"""
        if isinstance(item, dict):
            # Job title
            if item.get('title') or item.get('position'):
                title = item.get('title') or item.get('position')
                self.doc.add_paragraph(title, style='Job Title')
            
            # Company and dates
            company_date_parts = []
            if item.get('company') or item.get('organization'):
                company_date_parts.append(item.get('company') or item.get('organization'))
            if item.get('dates') or item.get('duration'):
                company_date_parts.append(item.get('dates') or item.get('duration'))
            
            if company_date_parts:
                self.doc.add_paragraph(' | '.join(company_date_parts), style='Company Date')
            
            # Responsibilities/achievements
            if item.get('responsibilities'):
                for resp in item['responsibilities']:
                    self.doc.add_paragraph(f"• {resp}", style='Bullet Point')
            elif item.get('achievements'):
                for achievement in item['achievements']:
                    self.doc.add_paragraph(f"• {achievement}", style='Bullet Point')
            elif item.get('description'):
                self.doc.add_paragraph(f"• {item['description']}", style='Bullet Point')
        else:
            self.doc.add_paragraph(f"• {item}", style='Bullet Point')
    
    def generate_resume(self, resume_data):
        """Generate a complete resume from structured data"""
        
        # Extract name and contact info
        name = "Professional Resume"
        contact_info = ""
        
        if isinstance(resume_data, dict):
            # Handle structured JSON data
            name = resume_data.get('name', resume_data.get('full_name', 'Professional Resume'))
            contact_info = resume_data.get('contact', resume_data.get('contact_info', ''))
            
            # Add header
            self.add_header(name, contact_info)
            
            # Define section order
            section_order = [
                'summary', 'professional_summary', 'objective',
                'experience', 'work_experience', 'professional_experience',
                'education', 'skills', 'technical_skills',
                'certifications', 'projects', 'achievements', 'awards'
            ]
            
            # Add sections in order
            for section_key in section_order:
                if section_key in resume_data:
                    section_title = section_key.replace('_', ' ').title()
                    self.add_section(section_title, resume_data[section_key])
            
            # Add any remaining sections
            for key, value in resume_data.items():
                if key not in ['name', 'full_name', 'contact', 'contact_info'] + section_order:
                    section_title = key.replace('_', ' ').title()
                    self.add_section(section_title, value)
        
        else:
            # Handle plain text - try to parse it
            parsed_data = self.parse_resume_content(str(resume_data))
            
            # Extract name from first line or use default
            lines = str(resume_data).split('\n')
            if lines:
                first_line = lines[0].strip()
                if len(first_line) < 50 and not first_line.isupper():
                    name = first_line
            
            self.add_header(name, contact_info)
            
            # Add parsed sections
            for section_title, content in parsed_data.items():
                self.add_section(section_title, content)
    
    def save_to_bytes(self):
        """Save the document to bytes for S3 upload"""
        import io
        doc_bytes = io.BytesIO()
        self.doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

def create_resume_docx(resume_content):
    """Main function to create a Word document from resume content"""
    template = ResumeTemplate()
    template.generate_resume(resume_content)
    return template.save_to_bytes()
