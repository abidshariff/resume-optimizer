"""
ATS-Friendly Word document generation using python-docx library.
Creates a clean, compact resume optimized for Applicant Tracking Systems with blue underlined headers.
"""

import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_enhanced_word_resume(resume_json):
    """Create compact, ATS-friendly Word document"""
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Tight margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Name - centered
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(resume_json.get('full_name', '').upper())
        name_run.font.size = Pt(16)
        name_run.bold = True
        name_p.space_after = Pt(3)
        
        # Contact - centered
        contact_parts = []
        if resume_json.get('email'): contact_parts.append(resume_json['email'])
        if resume_json.get('phone'): contact_parts.append(resume_json['phone'])
        if resume_json.get('location'): contact_parts.append(resume_json['location'])
        
        if contact_parts:
            contact_p = doc.add_paragraph()
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_p.add_run(' • '.join(contact_parts))
            contact_run.font.size = Pt(11)
            contact_p.space_after = Pt(8)
        
        # Professional Summary
        if resume_json.get('professional_summary'):
            summary_header = doc.add_paragraph()
            summary_run = summary_header.add_run('PROFESSIONAL SUMMARY')
            summary_run.font.size = Pt(12)
            summary_run.bold = True
            summary_header.space_before = Pt(6)
            summary_header.space_after = Pt(3)
            
            summary_p = doc.add_paragraph()
            summary_text = summary_p.add_run(resume_json['professional_summary'])
            summary_text.font.size = Pt(11)
            summary_p.space_after = Pt(6)
        
        # Core Competencies
        if resume_json.get('skills'):
            skills_header = doc.add_paragraph()
            skills_run = skills_header.add_run('CORE COMPETENCIES')
            skills_run.font.size = Pt(12)
            skills_run.bold = True
            skills_header.space_before = Pt(6)
            skills_header.space_after = Pt(3)
            
            skills = resume_json['skills']
            if isinstance(skills, list):
                skills_text = ' • '.join(skills)
            else:
                skills_text = str(skills)
            
            skills_p = doc.add_paragraph()
            skills_content = skills_p.add_run(skills_text)
            skills_content.font.size = Pt(11)
            skills_p.space_after = Pt(6)
        
        # Professional Experience
        if resume_json.get('experience'):
            exp_header = doc.add_paragraph()
            exp_run = exp_header.add_run('PROFESSIONAL EXPERIENCE')
            exp_run.font.size = Pt(12)
            exp_run.bold = True
            exp_header.space_before = Pt(6)
            exp_header.space_after = Pt(3)
            
            for exp in resume_json['experience']:
                # Job title
                job_p = doc.add_paragraph()
                job_run = job_p.add_run(exp.get('title', ''))
                job_run.font.size = Pt(11)
                job_run.bold = True
                job_p.space_before = Pt(4)
                job_p.space_after = Pt(1)
                
                # Company and dates
                company_p = doc.add_paragraph()
                company_text = f"{exp.get('company', '')} | {exp.get('dates', '')}"
                company_run = company_p.add_run(company_text)
                company_run.font.size = Pt(11)
                company_p.space_after = Pt(2)
                
                # Achievements
                for achievement in exp.get('achievements', []):
                    bullet_p = doc.add_paragraph()
                    bullet_p.style = 'List Bullet'
                    bullet_run = bullet_p.add_run(achievement)
                    bullet_run.font.size = Pt(10)
                    bullet_p.space_after = Pt(1)
        
        # Education
        if resume_json.get('education'):
            edu_header = doc.add_paragraph()
            edu_run = edu_header.add_run('EDUCATION')
            edu_run.font.size = Pt(12)
            edu_run.bold = True
            edu_header.space_before = Pt(6)
            edu_header.space_after = Pt(3)
            
            for edu in resume_json['education']:
                degree_p = doc.add_paragraph()
                degree_run = degree_p.add_run(edu.get('degree', ''))
                degree_run.font.size = Pt(11)
                degree_run.bold = True
                degree_p.space_after = Pt(1)
                
                school_p = doc.add_paragraph()
                school_text = f"{edu.get('institution', '')} | {edu.get('year', '')}"
                school_run = school_p.add_run(school_text)
                school_run.font.size = Pt(11)
                school_p.space_after = Pt(3)
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None
        contact_font = contact_style.font
        contact_font.name = 'Arial'
        contact_font.size = Pt(9)  # Reduced from 10
        contact_style.paragraph_format.space_after = Pt(8)
        contact_style.paragraph_format.space_before = Pt(0)
        
        # Section heading style - blue underlined
        heading_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(11)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)  # Light blue to match PDF
        heading_font.underline = True  # Add underline
        heading_style.paragraph_format.space_before = Pt(8)
        heading_style.paragraph_format.space_after = Pt(3)
        
        # Job title style - ultra compact
        job_title_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
        job_title_font = job_title_style.font
        job_title_font.name = 'Arial'
        job_title_font.size = Pt(10)
        job_title_font.bold = True
        job_title_style.paragraph_format.space_after = Pt(0)  # No spacing
        job_title_style.paragraph_format.space_before = Pt(4)  # Small gap between jobs only
        
        # Company style - ultra compact
        company_style = styles.add_style('Company', WD_STYLE_TYPE.PARAGRAPH)
        company_font = company_style.font
        company_font.name = 'Arial'
        company_font.size = Pt(10)
        company_style.paragraph_format.space_after = Pt(1)  # Minimal spacing before bullets
        company_style.paragraph_format.space_before = Pt(0)
        
        # Body text style - compact
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Arial'
        body_font.size = Pt(10)  # Reduced from 11
        body_style.paragraph_format.space_after = Pt(3)  # Reduced spacing
        body_style.paragraph_format.space_before = Pt(0)
        
        # Bullet style - ultra compact
        bullet_style = styles.add_style('BulletText', WD_STYLE_TYPE.PARAGRAPH)
        bullet_font = bullet_style.font
        bullet_font.name = 'Arial'
        bullet_font.size = Pt(10)
        bullet_style.paragraph_format.space_after = Pt(0)  # No spacing between bullets
        bullet_style.paragraph_format.space_before = Pt(0)
        bullet_style.paragraph_format.left_indent = Inches(0.15)  # Minimal indent
        bullet_style.paragraph_format.line_spacing = 1.0  # Single line spacing
        
        # Add name - center aligned
        name_para = doc.add_paragraph(resume_json.get('full_name', 'Name Not Provided'))
        name_para.style = name_style
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info - center aligned, cleaned
        contact_info = resume_json.get('contact_info', 'Contact information not provided')
        # Remove LinkedIn references
        contact_info = contact_info.replace('LinkedIn:', '').replace('linkedin.com/', '').replace('LinkedIn', '').strip()
        # Remove special characters that might appear as rhombus/squares, but keep bullets
        contact_info = re.sub(r'[^\w\s@.,|()\-+•·]', '', contact_info)
        # Clean up extra spaces and separators, then add bullet separators
        contact_parts = [part.strip() for part in contact_info.split('|') if part.strip() and 'linkedin' not in part.lower()]
        
        # Create contact paragraph with proper bullet formatting
        contact_para = doc.add_paragraph()
        contact_para.style = contact_style
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add each contact part with bullet separators
        for i, part in enumerate(contact_parts):
            if i > 0:
                # Add bullet separator as a run with larger size
                bullet_run = contact_para.add_run(' • ')
                bullet_run.font.name = 'Arial'
                bullet_run.font.size = Pt(11)  # Larger bullet size
                bullet_run.font.bold = True    # Make bullets bold for better visibility
            
            # Add the contact part
            part_run = contact_para.add_run(part)
            part_run.font.name = 'Arial'
            part_run.font.size = Pt(9)
        
        # Add professional summary - compact
        if resume_json.get('professional_summary'):
            summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
            summary_heading.style = heading_style
            
            summary_para = doc.add_paragraph(resume_json['professional_summary'])
            summary_para.style = body_style
        
        # Add core competencies/skills - compact format
        if resume_json.get('skills'):
            skills_heading = doc.add_paragraph('CORE COMPETENCIES')
            skills_heading.style = heading_style
            
            # Format skills in a single paragraph for compactness
            skills_list = resume_json['skills']
            skills_text = ' • '.join(skills_list[:12])  # Limit to 12 skills for space
            skills_para = doc.add_paragraph(skills_text)
            skills_para.style = body_style
        
        # Add professional experience - compact format
        if resume_json.get('experience'):
            exp_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE')
            exp_heading.style = heading_style
            
            for i, exp in enumerate(resume_json['experience']):
                # Job title - compact
                title_para = doc.add_paragraph(exp.get('title', 'Position Title'))
                title_para.style = job_title_style
                
                # Company and dates - same line for space efficiency
                company_name = exp.get('company', 'Company Name')
                dates = exp.get('dates', 'Dates')
                company_para = doc.add_paragraph(f"{company_name} | {dates}")
                company_para.style = company_style
                
                # Achievements - compact bullets
                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        bullet_para = doc.add_paragraph(f"• {achievement}")
                        bullet_para.style = bullet_style
        
        # Add education - compact format
        if resume_json.get('education'):
            edu_heading = doc.add_paragraph('EDUCATION')
            edu_heading.style = heading_style
            
            for edu in resume_json['education']:
                # Degree - compact
                degree = edu.get('degree', 'Degree')
                degree_para = doc.add_paragraph(degree)
                degree_para.style = job_title_style
                
                # Institution and details - single line
                institution = edu.get('institution', 'Institution')
                edu_details = institution
                
                if edu.get('dates'):
                    edu_details += f" | {edu['dates']}"
                if edu.get('gpa'):
                    edu_details += f" | GPA: {edu['gpa']}"
                
                edu_para = doc.add_paragraph(edu_details)
                edu_para.style = company_style
                
                # Additional details if brief
                if edu.get('details') and len(edu['details']) < 100:
                    details_para = doc.add_paragraph(f"• {edu['details']}")
                    details_para.style = bullet_style
        
        # Add certifications if available - compact
        if resume_json.get('certifications'):
            cert_heading = doc.add_paragraph('CERTIFICATIONS')
            cert_heading.style = heading_style
            
            for cert in resume_json['certifications']:
                cert_para = doc.add_paragraph(f"• {cert}")
                cert_para.style = bullet_style
        
        # Add projects if available - compact
        if resume_json.get('projects'):
            projects_heading = doc.add_paragraph('KEY PROJECTS')
            projects_heading.style = heading_style
            
            for project in resume_json['projects']:
                if isinstance(project, dict):
                    project_name = project.get('name', 'Project')
                    project_para = doc.add_paragraph(project_name)
                    project_para.style = job_title_style
                    
                    if project.get('description'):
                        desc_para = doc.add_paragraph(f"• {project['description']}")
                        desc_para.style = bullet_style
                else:
                    # If project is just a string
                    project_para = doc.add_paragraph(f"• {project}")
                    project_para.style = bullet_style
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        print(f"Error in enhanced Word generation: {str(e)}")
        # Fallback to basic generation if enhanced fails
        return create_basic_word_resume(resume_json)

def create_basic_word_resume(resume_json):
    """
    Fallback basic ATS-friendly Word document generation with compact formatting.
    """
    try:
        doc = Document()
        
        # Set very tight margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # Name - center aligned and clean
        name_para = doc.add_paragraph(resume_json.get('full_name', 'Name Not Provided'))
        name_run = name_para.runs[0]
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(14)
        name_run.font.bold = True
        name_para.paragraph_format.space_after = Pt(3)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info - center aligned, cleaned
        contact_info = resume_json.get('contact_info', 'Contact information not provided')
        # Remove LinkedIn references and special characters
        contact_info = contact_info.replace('LinkedIn:', '').replace('linkedin.com/', '').replace('LinkedIn', '').strip()
        contact_info = re.sub(r'[^\w\s@.,|()\-+•·]', '', contact_info)
        contact_parts = [part.strip() for part in contact_info.split('|') if part.strip() and 'linkedin' not in part.lower()]
        
        # Create contact paragraph with proper bullet formatting
        contact_para = doc.add_paragraph()
        contact_run = contact_para.runs[0] if contact_para.runs else contact_para.add_run('')
        contact_run.font.name = 'Arial'
        contact_run.font.size = Pt(9)  # Reduced size
        contact_para.paragraph_format.space_after = Pt(8)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add each contact part with bullet separators
        for i, part in enumerate(contact_parts):
            if i > 0:
                # Add bullet separator as a run with larger size
                bullet_run = contact_para.add_run(' • ')
                bullet_run.font.name = 'Arial'
                bullet_run.font.size = Pt(11)  # Larger bullet size
                bullet_run.font.bold = True    # Make bullets bold for better visibility
            
            # Add the contact part
            part_run = contact_para.add_run(part)
            part_run.font.name = 'Arial'
            part_run.font.size = Pt(9)
        
        # Professional Summary - blue underlined
        if resume_json.get('professional_summary'):
            summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
            summary_run = summary_heading.runs[0]
            summary_run.font.name = 'Arial'
            summary_run.font.bold = True
            summary_run.font.size = Pt(11)
            summary_run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)  # Light blue to match PDF
            summary_run.font.underline = True  # Underline
            summary_heading.paragraph_format.space_before = Pt(8)
            summary_heading.paragraph_format.space_after = Pt(3)
            
            summary_para = doc.add_paragraph(resume_json['professional_summary'])
            summary_para.runs[0].font.name = 'Arial'
            summary_para.runs[0].font.size = Pt(9)  # Reduced size
            summary_para.paragraph_format.space_after = Pt(3)
            summary_para.paragraph_format.space_after = Pt(3)
        
        # Skills - blue underlined
        if resume_json.get('skills'):
            skills_heading = doc.add_paragraph('CORE COMPETENCIES')
            skills_run = skills_heading.runs[0]
            skills_run.font.name = 'Arial'
            skills_run.font.bold = True
            skills_run.font.size = Pt(11)
            skills_run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)  # Light blue to match PDF
            skills_run.font.underline = True  # Underline
            skills_heading.paragraph_format.space_before = Pt(8)
            skills_heading.paragraph_format.space_after = Pt(3)
            
            # Format skills compactly
            skills_text = ' • '.join(resume_json['skills'][:12])
            skills_para = doc.add_paragraph(skills_text)
            skills_para.runs[0].font.name = 'Arial'
            skills_para.runs[0].font.size = Pt(9)  # Reduced size
            skills_para.paragraph_format.space_after = Pt(3)
            skills_para.paragraph_format.space_after = Pt(3)
        
        # Experience - blue underlined
        if resume_json.get('experience'):
            exp_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE')
            exp_run = exp_heading.runs[0]
            exp_run.font.name = 'Arial'
            exp_run.font.bold = True
            exp_run.font.size = Pt(11)
            exp_run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)  # Light blue to match PDF
            exp_run.font.underline = True  # Underline
            exp_heading.paragraph_format.space_before = Pt(8)
            exp_heading.paragraph_format.space_after = Pt(3)
            
            for exp in resume_json['experience']:
                # Job title - ultra compact
                title_para = doc.add_paragraph(exp.get('title', 'Position Title'))
                title_run = title_para.runs[0]
                title_run.font.name = 'Arial'
                title_run.font.bold = True
                title_run.font.size = Pt(10)
                title_para.paragraph_format.space_before = Pt(4)  # Small gap between jobs
                title_para.paragraph_format.space_after = Pt(0)   # No spacing
                
                # Company and dates - ultra compact
                company_para = doc.add_paragraph(f"{exp.get('company', 'Company')} | {exp.get('dates', 'Dates')}")
                company_run = company_para.runs[0]
                company_run.font.name = 'Arial'
                company_run.font.size = Pt(10)
                company_para.paragraph_format.space_after = Pt(1)  # Minimal space before bullets
                
                # Achievements - ultra compact bullets
                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        bullet_para = doc.add_paragraph(f"• {achievement}")
                        bullet_run = bullet_para.runs[0]
                        bullet_run.font.name = 'Arial'
                        bullet_run.font.size = Pt(10)
                        bullet_para.paragraph_format.space_after = Pt(0)  # No spacing between bullets
                        bullet_para.paragraph_format.space_before = Pt(0)
                        bullet_para.paragraph_format.left_indent = Inches(0.15)
                        bullet_para.paragraph_format.line_spacing = 1.0
        
        # Education - blue underlined
        if resume_json.get('education'):
            edu_heading = doc.add_paragraph('EDUCATION')
            edu_run = edu_heading.runs[0]
            edu_run.font.name = 'Arial'
            edu_run.font.bold = True
            edu_run.font.size = Pt(11)
            edu_run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)  # Light blue to match PDF
            edu_run.font.underline = True  # Underline
            edu_heading.paragraph_format.space_before = Pt(8)
            edu_heading.paragraph_format.space_after = Pt(3)
            
            for edu in resume_json['education']:
                degree_para = doc.add_paragraph(edu.get('degree', 'Degree'))
                degree_run = degree_para.runs[0]
                degree_run.font.name = 'Arial'
                degree_run.font.bold = True
                degree_run.font.size = Pt(10)
                degree_para.paragraph_format.space_before = Pt(3)
                degree_para.paragraph_format.space_after = Pt(1)
                
                edu_info = edu.get('institution', 'Institution')
                if edu.get('dates'):
                    edu_info += f" | {edu['dates']}"
                
                edu_para = doc.add_paragraph(edu_info)
                edu_para.runs[0].font.name = 'Arial'
                edu_para.runs[0].font.size = Pt(10)
                edu_para.paragraph_format.space_after = Pt(2)
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        print(f"Error in basic Word generation: {str(e)}")
        raise e
