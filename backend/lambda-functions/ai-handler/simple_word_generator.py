"""
Simple Word document generation without lxml dependency.
Uses only python-docx which has fewer platform-specific dependencies.
"""

import io
import boto3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_resume_simple(resume_json, template_bucket=None, template_key=None):
    """
    Generate a Word document resume without using docxtpl.
    Creates the document programmatically using python-docx only.
    
    Args:
        resume_json (dict): Resume data
        template_bucket (str): Not used in this implementation
        template_key (str): Not used in this implementation
    
    Returns:
        bytes: Generated Word document as bytes
    """
    
    try:
        # Create a new document
        doc = Document()
        
        # Set narrow margins for one-page layout
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # === HEADER SECTION ===
        # Name - Large, bold, centered
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(resume_json.get('full_name', 'Your Name'))
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
        
        # Contact info - centered, smaller
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(resume_json.get('contact_info', 'Your Contact Information'))
        contact_run.font.name = 'Calibri'
        contact_run.font.size = Pt(11)
        
        # Add some space
        doc.add_paragraph()
        
        # === PROFESSIONAL SUMMARY ===
        summary_header = doc.add_paragraph()
        summary_header_run = summary_header.add_run('PROFESSIONAL SUMMARY')
        summary_header_run.font.name = 'Calibri'
        summary_header_run.font.size = Pt(12)
        summary_header_run.font.bold = True
        summary_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
        summary_header_run.underline = True
        
        summary_para = doc.add_paragraph(resume_json.get('professional_summary', 'Professional summary not available'))
        summary_para.style.font.name = 'Calibri'
        summary_para.style.font.size = Pt(11)
        
        # === CORE COMPETENCIES ===
        skills_header = doc.add_paragraph()
        skills_header_run = skills_header.add_run('CORE COMPETENCIES')
        skills_header_run.font.name = 'Calibri'
        skills_header_run.font.size = Pt(12)
        skills_header_run.font.bold = True
        skills_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
        skills_header_run.underline = True
        
        # Skills as bullet points
        skills = resume_json.get('skills', [])
        if skills:
            skills_text = ' • '.join(skills)
            skills_para = doc.add_paragraph()
            skills_run = skills_para.add_run(f"• {skills_text}")
            skills_run.font.name = 'Calibri'
            skills_run.font.size = Pt(11)
        
        # === PROFESSIONAL EXPERIENCE ===
        exp_header = doc.add_paragraph()
        exp_header_run = exp_header.add_run('PROFESSIONAL EXPERIENCE')
        exp_header_run.font.name = 'Calibri'
        exp_header_run.font.size = Pt(12)
        exp_header_run.font.bold = True
        exp_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
        exp_header_run.underline = True
        
        # Experience entries
        experience = resume_json.get('experience', [])
        for job in experience:
            # Job title - bold
            job_title_para = doc.add_paragraph()
            job_title_run = job_title_para.add_run(job.get('title', 'Job Title'))
            job_title_run.font.name = 'Calibri'
            job_title_run.font.size = Pt(11)
            job_title_run.font.bold = True
            
            # Company and dates - italic
            company_para = doc.add_paragraph()
            company_text = f"{job.get('company', 'Company')} | {job.get('dates', 'Dates')}"
            company_run = company_para.add_run(company_text)
            company_run.font.name = 'Calibri'
            company_run.font.size = Pt(10)
            company_run.font.italic = True
            
            # Responsibilities
            responsibilities = job.get('responsibilities', [])
            if isinstance(responsibilities, str):
                responsibilities = [responsibilities]
            
            for responsibility in responsibilities:
                resp_para = doc.add_paragraph()
                resp_run = resp_para.add_run(f"• {responsibility}")
                resp_run.font.name = 'Calibri'
                resp_run.font.size = Pt(10)
            
            # Add space between jobs
            doc.add_paragraph()
        
        # === EDUCATION ===
        edu_header = doc.add_paragraph()
        edu_header_run = edu_header.add_run('EDUCATION')
        edu_header_run.font.name = 'Calibri'
        edu_header_run.font.size = Pt(12)
        edu_header_run.font.bold = True
        edu_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
        edu_header_run.underline = True
        
        # Education entries
        education = resume_json.get('education', [])
        for edu in education:
            # Degree - bold
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.get('degree', 'Degree'))
            degree_run.font.name = 'Calibri'
            degree_run.font.size = Pt(11)
            degree_run.font.bold = True
            
            # Institution and dates
            institution_para = doc.add_paragraph()
            institution_text = f"{edu.get('institution', 'Institution')} | {edu.get('dates', 'Dates')}"
            institution_run = institution_para.add_run(institution_text)
            institution_run.font.name = 'Calibri'
            institution_run.font.size = Pt(10)
            institution_run.font.italic = True
        
        # Save to bytes
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        
        return output_stream.getvalue()
        
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        raise e
