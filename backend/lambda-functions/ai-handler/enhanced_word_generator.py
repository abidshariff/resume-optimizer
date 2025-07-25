"""
Enhanced Word document generation using python-docx library.
Creates a professionally formatted resume in .docx format with page length control.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_enhanced_word_resume(resume_json):
    """
    Create a professionally formatted Word document using python-docx with optimized spacing.
    
    Args:
        resume_json (dict): Resume data
    
    Returns:
        bytes: Generated Word document as bytes
    """
    
    try:
        # Create a new Document
        doc = Document()
        
        # Set document margins for better space utilization
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Define styles with optimized spacing
        styles = doc.styles
        
        # Name style - slightly smaller for space efficiency
        name_style = styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
        name_font = name_style.font
        name_font.name = 'Calibri'
        name_font.size = Pt(16)  # Reduced from 18
        name_font.bold = True
        name_font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dark blue
        name_style.paragraph_format.space_after = Pt(3)  # Reduced spacing
        
        # Contact style - more compact
        contact_style = styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
        contact_font = contact_style.font
        contact_font.name = 'Calibri'
        contact_font.size = Pt(10)  # Reduced from 11
        contact_font.color.rgb = RGBColor(0x4F, 0x4F, 0x4F)  # Dark gray
        contact_style.paragraph_format.space_after = Pt(6)  # Reduced spacing
        
        # Section heading style - more compact
        heading_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.size = Pt(12)  # Reduced from 14
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dark blue
        heading_style.paragraph_format.space_before = Pt(8)  # Reduced spacing
        heading_style.paragraph_format.space_after = Pt(4)   # Reduced spacing
        
        # Job title style - more compact
        job_title_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
        job_title_font = job_title_style.font
        job_title_font.name = 'Calibri'
        job_title_font.size = Pt(11)  # Reduced from 12
        job_title_font.bold = True
        job_title_style.paragraph_format.space_after = Pt(2)  # Reduced spacing
        
        # Company style - more compact
        company_style = styles.add_style('Company', WD_STYLE_TYPE.PARAGRAPH)
        company_font = company_style.font
        company_font.name = 'Calibri'
        company_font.size = Pt(10)  # Reduced from 11
        company_font.italic = True
        company_style.paragraph_format.space_after = Pt(3)  # Reduced spacing
        
        # Body text style - more compact
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(10)  # Reduced from 11
        body_style.paragraph_format.space_after = Pt(2)  # Reduced spacing
        body_style.paragraph_format.left_indent = Inches(0.2)
        
        # Bullet style - more compact
        bullet_style = styles.add_style('BulletText', WD_STYLE_TYPE.PARAGRAPH)
        bullet_font = bullet_style.font
        bullet_font.name = 'Calibri'
        bullet_font.size = Pt(10)  # Reduced from 11
        bullet_style.paragraph_format.space_after = Pt(1)  # Minimal spacing
        bullet_style.paragraph_format.left_indent = Inches(0.3)
        
        # Add name
        name_para = doc.add_paragraph(resume_json.get('full_name', 'Name Not Provided'))
        name_para.style = name_style
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info
        contact_para = doc.add_paragraph(resume_json.get('contact_info', 'Contact information not provided'))
        contact_para.style = contact_style
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add professional summary
        if resume_json.get('professional_summary'):
            summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
            summary_heading.style = heading_style
            
            summary_para = doc.add_paragraph(resume_json['professional_summary'])
            summary_para.style = body_style
        
        # Add skills section - format as comma-separated for space efficiency
        if resume_json.get('skills'):
            skills_heading = doc.add_paragraph('CORE COMPETENCIES')
            skills_heading.style = heading_style
            
            # Format skills as comma-separated text for better space utilization
            skills_text = ' • '.join(resume_json['skills'][:12])  # Limit to 12 skills
            skills_para = doc.add_paragraph(skills_text)
            skills_para.style = body_style
        
        # Add experience section
        if resume_json.get('experience'):
            exp_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE')
            exp_heading.style = heading_style
            
            for exp in resume_json['experience']:
                # Job title
                title_para = doc.add_paragraph(exp.get('title', 'Position Title'))
                title_para.style = job_title_style
                
                # Company and dates on same line for space efficiency
                company_info = f"{exp.get('company', 'Company Name')} | {exp.get('dates', 'Dates')}"
                company_para = doc.add_paragraph(company_info)
                company_para.style = company_style
                
                # Achievements - limit bullet points for conciseness
                if exp.get('achievements'):
                    for achievement in exp['achievements'][:4]:  # Limit to 4 bullets max
                        bullet_para = doc.add_paragraph(f"• {achievement}")
                        bullet_para.style = bullet_style
        
        # Add education section
        if resume_json.get('education'):
            edu_heading = doc.add_paragraph('EDUCATION')
            edu_heading.style = heading_style
            
            for edu in resume_json['education']:
                # Format education on fewer lines for space efficiency
                edu_text = f"{edu.get('degree', 'Degree')} | {edu.get('institution', 'Institution')}"
                if edu.get('dates'):
                    edu_text += f" | {edu['dates']}"
                
                edu_para = doc.add_paragraph(edu_text)
                edu_para.style = body_style
                
                # Add details only if they're brief and relevant
                if edu.get('details') and len(edu['details']) < 100:
                    details_para = doc.add_paragraph(edu['details'])
                    details_para.style = bullet_style
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        print(f"Error in enhanced Word generation: {str(e)}")
        # Fallback to basic generation if enhanced fails
        return create_basic_word_resume(resume_json)

def create_basic_word_resume(resume_json):
    """
    Fallback basic Word document generation.
    """
    try:
        doc = Document()
        
        # Basic formatting with minimal spacing
        doc.add_paragraph(resume_json.get('full_name', 'Name Not Provided')).runs[0].font.size = Pt(16)
        doc.add_paragraph(resume_json.get('contact_info', 'Contact information not provided')).runs[0].font.size = Pt(10)
        
        if resume_json.get('professional_summary'):
            doc.add_paragraph('PROFESSIONAL SUMMARY').runs[0].font.bold = True
            doc.add_paragraph(resume_json['professional_summary'])
        
        if resume_json.get('skills'):
            doc.add_paragraph('SKILLS').runs[0].font.bold = True
            doc.add_paragraph(' • '.join(resume_json['skills'][:12]))
        
        if resume_json.get('experience'):
            doc.add_paragraph('EXPERIENCE').runs[0].font.bold = True
            for exp in resume_json['experience']:
                doc.add_paragraph(f"{exp.get('title', '')} - {exp.get('company', '')}")
                doc.add_paragraph(exp.get('dates', ''))
                if exp.get('achievements'):
                    for achievement in exp['achievements'][:3]:
                        doc.add_paragraph(f"• {achievement}")
        
        if resume_json.get('education'):
            doc.add_paragraph('EDUCATION').runs[0].font.bold = True
            for edu in resume_json['education']:
                doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        print(f"Error in basic Word generation: {str(e)}")
        raise e
        contact_font.name = 'Calibri'
        contact_font.size = Pt(11)
        
        # Section header style
        header_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        header_font = header_style.font
        header_font.name = 'Calibri'
        header_font.size = Pt(14)
        header_font.bold = True
        header_font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dark blue
        
        # Job title style
        job_title_style = styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
        job_title_font = job_title_style.font
        job_title_font.name = 'Calibri'
        job_title_font.size = Pt(12)
        job_title_font.bold = True
        
        # Job details style
        job_details_style = styles.add_style('JobDetails', WD_STYLE_TYPE.PARAGRAPH)
        job_details_font = job_details_style.font
        job_details_font.name = 'Calibri'
        job_details_font.size = Pt(11)
        job_details_font.italic = True
        
        # Normal text style
        normal_style = styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
        # Add name
        name = resume_json.get('full_name', 'Your Name')
        name_paragraph = doc.add_paragraph(name, style='Name')
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info
        contact = resume_json.get('contact_info', 'Your Contact Information')
        contact_paragraph = doc.add_paragraph(contact, style='Contact')
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line
        add_horizontal_line(doc)
        
        # Add Professional Summary section
        doc.add_paragraph('PROFESSIONAL SUMMARY', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        summary = resume_json.get('professional_summary', 'Professional summary not available')
        doc.add_paragraph(summary, style='NormalText')
        
        # Add Skills section
        doc.add_paragraph('SKILLS', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        skills = resume_json.get('skills', [])
        if skills:
            skills_paragraph = doc.add_paragraph(style='NormalText')
            for i, skill in enumerate(skills):
                if i > 0:
                    skills_paragraph.add_run(' • ')
                skills_paragraph.add_run(skill)
        
        # Add Experience section
        doc.add_paragraph('EXPERIENCE', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        experience = resume_json.get('experience', [])
        for job in experience:
            # Job title
            title = job.get('title', 'Job Title')
            doc.add_paragraph(title, style='JobTitle')
            
            # Company and dates
            company = job.get('company', 'Company')
            dates = job.get('dates', 'Dates')
            doc.add_paragraph(f"{company} | {dates}", style='JobDetails')
            
            # Achievements
            achievements = job.get('achievements', [])
            for achievement in achievements:
                achievement_para = doc.add_paragraph(style='NormalText')
                achievement_para.add_run('• ').bold = True
                achievement_para.add_run(achievement)
            
            # Add space between jobs
            doc.add_paragraph()
        
        # Add Education section
        doc.add_paragraph('EDUCATION', style='SectionHeader')
        add_horizontal_line(doc, thin=True)
        education = resume_json.get('education', [])
        for edu in education:
            # Degree
            degree = edu.get('degree', 'Degree')
            doc.add_paragraph(degree, style='JobTitle')
            
            # Institution and dates
            institution = edu.get('institution', 'Institution')
            dates = edu.get('dates', 'Dates')
            doc.add_paragraph(f"{institution} | {dates}", style='JobDetails')
            
            # Details if available
            if 'details' in edu and edu['details']:
                doc.add_paragraph(edu['details'], style='NormalText')
        
        # Save the document to a bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error creating enhanced Word document: {str(e)}")
        raise e

def add_horizontal_line(doc, thin=False):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    
    # Create the line
    run = p.add_run()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6' if thin else '12')  # 6 for thin line, 12 for thick line
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '1F4E79')  # Dark blue
    border.append(bottom)
    
    # Add the border to the paragraph
    p._p.get_or_add_pPr().append(border)
