#!/usr/bin/env python3
"""
Create a professional one-page resume template optimized for docxtpl.
This version uses proper docxtpl syntax and formatting.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_professional_template():
    """Create a clean, professional resume template."""
    
    doc = Document()
    
    # Set narrow margins for one-page layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # === HEADER SECTION ===
    # Name - Large, bold, centered
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('{{full_name}}')
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)  # Professional dark blue
    
    # Contact info - centered, smaller
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run('{{contact_info}}')
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(11)
    
    # Add some space
    doc.add_paragraph()
    
    # === PROFESSIONAL SUMMARY ===
    summary_header = doc.add_paragraph()
    summary_header_run = summary_header.add_run('PROFESSIONAL SUMMARY')
    summary_header_run.font.name = 'Calibri'
    summary_header_run.font.size = Pt(12)
    summary_header_run.font.bold = True
    summary_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    summary_header_run.underline = True
    
    summary_para = doc.add_paragraph('{{professional_summary}}')
    summary_para.style.font.name = 'Calibri'
    summary_para.style.font.size = Pt(11)
    
    # === CORE COMPETENCIES ===
    skills_header = doc.add_paragraph()
    skills_header_run = skills_header.add_run('CORE COMPETENCIES')
    skills_header_run.font.name = 'Calibri'
    skills_header_run.font.size = Pt(12)
    skills_header_run.font.bold = True
    skills_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    skills_header_run.underline = True
    
    # Skills as bullet points
    skills_para = doc.add_paragraph()
    skills_run = skills_para.add_run('{%p for skill in skills %}• {{skill}}{%p endfor %}')
    skills_run.font.name = 'Calibri'
    skills_run.font.size = Pt(11)
    
    # === PROFESSIONAL EXPERIENCE ===
    exp_header = doc.add_paragraph()
    exp_header_run = exp_header.add_run('PROFESSIONAL EXPERIENCE')
    exp_header_run.font.name = 'Calibri'
    exp_header_run.font.size = Pt(12)
    exp_header_run.font.bold = True
    exp_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    exp_header_run.underline = True
    
    # Experience entries
    exp_para = doc.add_paragraph('{%p for job in experience %}')
    
    # Job title - bold
    job_title_para = doc.add_paragraph()
    job_title_run = job_title_para.add_run('{{job.title}}')
    job_title_run.font.name = 'Calibri'
    job_title_run.font.size = Pt(11)
    job_title_run.font.bold = True
    
    # Company and dates - italic
    company_para = doc.add_paragraph()
    company_run = company_para.add_run('{{job.company}} | {{job.dates}}')
    company_run.font.name = 'Calibri'
    company_run.font.size = Pt(10)
    company_run.font.italic = True
    
    # Responsibilities
    resp_para = doc.add_paragraph('{%p for responsibility in job.responsibilities %}• {{responsibility}}{%p endfor %}')
    resp_para.style.font.name = 'Calibri'
    resp_para.style.font.size = Pt(10)
    
    exp_end_para = doc.add_paragraph('{%p endfor %}')
    
    # === EDUCATION ===
    edu_header = doc.add_paragraph()
    edu_header_run = edu_header.add_run('EDUCATION')
    edu_header_run.font.name = 'Calibri'
    edu_header_run.font.size = Pt(12)
    edu_header_run.font.bold = True
    edu_header_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    edu_header_run.underline = True
    
    # Education entries
    edu_para = doc.add_paragraph('{%p for edu in education %}')
    
    # Degree - bold
    degree_para = doc.add_paragraph()
    degree_run = degree_para.add_run('{{edu.degree}}')
    degree_run.font.name = 'Calibri'
    degree_run.font.size = Pt(11)
    degree_run.font.bold = True
    
    # Institution and dates
    institution_para = doc.add_paragraph()
    institution_run = institution_para.add_run('{{edu.institution}} | {{edu.dates}}')
    institution_run.font.name = 'Calibri'
    institution_run.font.size = Pt(10)
    institution_run.font.italic = True
    
    edu_end_para = doc.add_paragraph('{%p endfor %}')
    
    return doc

def main():
    """Create and save the resume template."""
    doc = create_professional_template()
    doc.save('professional_resume_template.docx')
    print("Professional resume template created: professional_resume_template.docx")
    print("\nThis template uses docxtpl syntax and includes:")
    print("- Professional formatting with consistent fonts and colors")
    print("- One-page optimized layout with narrow margins")
    print("- Section headers with underlines and professional blue color")
    print("- Proper spacing and hierarchy")
    print("- Support for dynamic content with loops and conditionals")

if __name__ == '__main__':
    main()
